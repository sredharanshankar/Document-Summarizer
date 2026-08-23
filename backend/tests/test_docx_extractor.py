import pytest

from app.services.docx_extractor import extract_docx_text
from app.utils.errors import CorruptedFileError


def test_extracts_paragraph_text(sample_docx_bytes: bytes) -> None:
    text = extract_docx_text(sample_docx_bytes)
    assert "sample DOCX document" in text


def test_skips_blank_paragraphs() -> None:
    import io

    import docx

    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("")  # blank paragraph, e.g. spacing in the original doc
    document.add_paragraph("Second paragraph.")
    buffer = io.BytesIO()
    document.save(buffer)

    text = extract_docx_text(buffer.getvalue())

    assert text == "First paragraph.\n\nSecond paragraph."


def test_rejects_non_docx_bytes() -> None:
    with pytest.raises(CorruptedFileError):
        extract_docx_text(b"this is not a docx file at all")
